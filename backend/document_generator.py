"""
Document generation utilities for PDF, DOCX, and Markdown
"""
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown
import os
from typing import Dict

class DocumentGenerator:
    @staticmethod
    def generate_markdown(book_data: Dict, output_path: str) -> str:
        """Generate Markdown file from book data"""
        try:
            md_content = []
            
            # Title Page
            md_content.append("# " + "📚 Bollywood Cloud Computing Book\n")
            md_content.append(book_data.get("title_page", ""))
            md_content.append("\n---\n")
            
            # Table of Contents
            md_content.append("## 📖 Table of Contents\n")
            md_content.append(book_data.get("toc", ""))
            md_content.append("\n---\n")
            
            # Chapters
            for chapter in book_data.get("chapters", []):
                md_content.append(f"\n## Chapter {chapter['number']}: {chapter['title']}\n")
                md_content.append(chapter["content"])
                md_content.append("\n---\n")
            
            # Write to file
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("\n".join(md_content))
            
            return output_path
        except Exception as e:
            raise Exception(f"Error generating Markdown: {str(e)}")
    
    @staticmethod
    def generate_docx(book_data: Dict, output_path: str) -> str:
        """Generate DOCX file from book data"""
        try:
            doc = Document()
            
            # Title Page
            title = doc.add_heading('📚 Bollywood Cloud Computing Book', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            title_content = doc.add_paragraph(book_data.get("title_page", ""))
            title_content.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_page_break()
            
            # Table of Contents
            doc.add_heading('📖 Table of Contents', level=1)
            doc.add_paragraph(book_data.get("toc", ""))
            doc.add_page_break()
            
            # Chapters
            for chapter in book_data.get("chapters", []):
                heading = doc.add_heading(
                    f"Chapter {chapter['number']}: {chapter['title']}", 
                    level=1
                )
                
                # Add chapter content
                # Split by pages and format
                content = chapter["content"]
                paragraphs = content.split('\n\n')
                
                for para in paragraphs:
                    if para.strip():
                        p = doc.add_paragraph(para.strip())
                        # Check if it's a heading
                        if para.startswith('Page ') or para.startswith('━'):
                            run = p.runs[0]
                            run.bold = True
                            run.font.size = Pt(12)
                
                doc.add_page_break()
            
            # Save document
            doc.save(output_path)
            return output_path
        except Exception as e:
            raise Exception(f"Error generating DOCX: {str(e)}")
    
    @staticmethod
    def generate_pdf(book_data: Dict, output_path: str) -> str:
        """Generate PDF file from book data"""
        try:
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18,
            )
            
            # Container for the 'Flowable' objects
            elements = []
            
            # Define styles
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor='#FF6B6B',
                spaceAfter=30,
                alignment=TA_CENTER,
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor='#4ECDC4',
                spaceAfter=12,
            )
            
            normal_style = styles['Normal']
            normal_style.fontSize = 11
            normal_style.leading = 14
            
            # Title Page
            elements.append(Paragraph("📚 Bollywood Cloud Computing Book", title_style))
            elements.append(Spacer(1, 0.2 * inch))
            
            title_text = book_data.get("title_page", "").replace('\n', '<br/>')
            elements.append(Paragraph(title_text, normal_style))
            elements.append(PageBreak())
            
            # Table of Contents
            elements.append(Paragraph("📖 Table of Contents", heading_style))
            toc_text = book_data.get("toc", "").replace('\n', '<br/>')
            elements.append(Paragraph(toc_text, normal_style))
            elements.append(PageBreak())
            
            # Chapters
            for chapter in book_data.get("chapters", []):
                chapter_title = f"Chapter {chapter['number']}: {chapter['title']}"
                elements.append(Paragraph(chapter_title, heading_style))
                elements.append(Spacer(1, 0.2 * inch))
                
                # Add chapter content
                content = chapter["content"].replace('\n\n', '<br/><br/>')
                content = content.replace('\n', '<br/>')
                
                # Handle special characters
                content = content.replace('&', '&amp;')
                content = content.replace('<', '&lt;').replace('>', '&gt;')
                content = content.replace('<br/>', '<br/>').replace('&lt;br/&gt;', '<br/>')
                
                elements.append(Paragraph(content, normal_style))
                elements.append(PageBreak())
            
            # Build PDF
            doc.build(elements)
            return output_path
        except Exception as e:
            raise Exception(f"Error generating PDF: {str(e)}")
